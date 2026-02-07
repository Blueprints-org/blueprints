"""Tests for the ReportToWordConverter class."""

import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument

from blueprints.utils._report_to_word import _ReportToWordConverter
from blueprints.utils.report import Report


class TestReportToWordConverter:
    """Tests for the ReportToWordConverter class."""

    def test_empty_input_returns_empty_document(self) -> None:
        """Test that an empty LaTeX string returns an empty Document."""
        assert _ReportToWordConverter("")

    def test_complex_document_conversion(self) -> None:
        """Test conversion of a complex LaTeX document."""
        report = Report("Testing Title")
        report.add_heading("Testing Section", level=1)
        report.add_heading("Testing Subsection", level=2)
        report.add_heading("Testing Subsubsection", level=3)
        report.add_paragraph("Bold and", bold=True).add_paragraph(" normal text.").add_newline(n=2)
        report.add_paragraph("And Italic text.", italic=True).add_paragraph(" And also bold and italic.", bold=True, italic=True).add_newline()
        report.add_paragraph("Here is an inline equation: $E=mc^2$ within the text.").add_newline()
        report.add_paragraph("test").add_equation("E=mc^2", tag="1", inline=True).add_paragraph("more text.").add_newline()
        report.add_paragraph("test").add_equation("E=mc^2", inline=True).add_paragraph("more text.").add_newline()
        report.add_equation("E=mc^2", tag="4")
        report.add_equation(r"\int_a^b f(x)dx = F(b) - F(a)")
        report.add_equation(
            "A_v = max(A - 2 \\cdot b \\cdot t_f + (t_w + 2 \\cdot r) \\cdot t_f;"
            " \\eta \\cdot h_w \\cdot t_w) = max("
            "12324.48 - 2 \\cdot 297.00 \\cdot 16.00 + (8.00 + 2 \\cdot 28.50) \\cdot 16.00;"
            " 1.00 \\cdot 265.00 \\cdot 8.00) = 3860.48 \\ mm^2 "
        )
        report.add_list(["One", ["A", "B", "C"], "Two", ["A", ["I", "II", ["A", "B"], "III"]]], style="numbered")
        report.add_list(["First", "Second", ["Subfirst", "Subsecond"], "Third"], style="bulleted")
        report.add_paragraph("Here is a table:")
        report.add_table(
            headers=["Header 1", "Header 2", "Header 3 with math $E=mc^2$"],
            rows=[["Row 1 Col 1", "Row 1 Col 2 with inline math $a^2 + b^2 = c^2$", "Row 1 Col 3"], ["Row 2 Col 1", "Row 2 Col 2", "Row 2 Col 3"]],
        )
        report.add_heading("Another Section", level=1)
        report.add_heading("Another Subsection", level=2)
        report.add_heading("Yet another Subsection", level=2)
        report.add_heading("Yet another Subsubsection", level=3)
        report.add_heading("Yet another Subsubsection v2", level=3)

        # Use cross-platform path that works on both Windows and Linux/Mac
        logo_path = Path("docs") / "_overrides" / "assets" / "images" / "logo-light-mode.png"
        report.add_figure(str(logo_path), width=0.4, caption="Description of the image.")
        latex_content = report.to_latex(language="zh")
        assert _ReportToWordConverter(latex_content).document

    def test_indented_enumerate_above_26(self) -> None:
        """Test that enumerate items above 26 are handled correctly."""
        report = Report("Testing Title")
        report.add_list(["One", [f"Item {i + 1}" for i in range(30)], "Two"], style="numbered")
        latex_content = report.to_latex()
        converter = _ReportToWordConverter(latex_content)
        document = converter.document
        assert document

    def test_to_word_returns_bytes_when_path_is_none(self) -> None:
        """Test that to_word() returns bytes when path is None."""
        report = Report(title="Test Report")
        report.add_heading("Introduction")
        report.add_paragraph("Some test content")

        # Call without path argument
        result = report.to_word()

        # Verify result is bytes
        assert isinstance(result, bytes)
        assert len(result) > 0

        # Verify bytes are valid DOCX format
        docx_buffer = BytesIO(result)
        docx_doc = DocxDocument(docx_buffer)
        assert docx_doc

    def test_to_word_with_bytesio_object(self) -> None:
        """Test that to_word() works with BytesIO objects."""
        report = Report(title="Test Report")
        report.add_heading("Introduction")
        report.add_paragraph("Some test content")

        # Call with BytesIO object
        buffer = BytesIO()
        result = report.to_word(buffer)

        # Verify return value is None
        assert result is None

        # Verify buffer contains valid DOCX data
        buffer.seek(0)
        assert len(buffer.getvalue()) > 0

        # Verify bytes are valid DOCX format
        docx_doc = DocxDocument(buffer)
        assert docx_doc

    def test_to_word_with_file_path_string(self) -> None:
        """Test that to_word() saves to file when given a string path."""
        report = Report(title="Test Report")
        report.add_heading("Introduction")
        report.add_paragraph("Some test content")

        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = str(Path(tmpdir) / "report.docx")

            # Call with file path
            result = report.to_word(file_path)

            # Verify return value is None
            assert result is None

            # Verify file was created
            assert Path(file_path).exists()

            # Verify file contains valid DOCX
            with open(file_path, "rb") as f:
                docx_doc = DocxDocument(f)
                assert docx_doc

    def test_to_word_with_pathlib_path(self) -> None:
        """Test that to_word() saves to file when given a Path object."""
        report = Report(title="Test Report")
        report.add_heading("Introduction")
        report.add_paragraph("Some test content")

        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "report.docx"

            # Call with Path object
            result = report.to_word(file_path)

            # Verify return value is None
            assert result is None

            # Verify file was created
            assert file_path.exists()

            # Verify file contains valid DOCX
            docx_doc = DocxDocument(str(file_path))
            assert docx_doc

    def test_to_word_bytes_with_language_translation(self) -> None:
        """Test that to_word() returns bytes with translated content."""
        report = Report(title="Test Report")
        report.add_heading("Introduction")
        report.add_paragraph("Some test content")

        # Call without path but with language translation
        result = report.to_word(language="zh")

        # Verify result is bytes
        assert isinstance(result, bytes)
        assert len(result) > 0

        # Verify bytes are valid DOCX format
        docx_buffer = BytesIO(result)
        docx_doc = DocxDocument(docx_buffer)
        assert docx_doc

    def test_empty_report_title_not_added(self) -> None:
        """Test that an empty report title is not added to the Word document."""
        report = Report()
        report.add_paragraph("Some content without a title.")

        # Call without path argument
        result = report.to_word()

        # Verify result is bytes
        assert isinstance(result, bytes)
        assert len(result) > 0
