"""Module to convert LaTeX (with text and equations) to a Word document."""

import re

import latex2mathml.converter
import mathml2omml
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import parse_xml
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Pt, RGBColor


class LatexToWordConverter:
    r"""
    Limitations and Usage Notes:
    - This converter currently supports LaTeX input containing text and equations only. Figures and tables are not yet supported.
    - Use \newline to indicate new lines. Do not use the double backslash \\ for line breaks.
    - All text must be wrapped in \text{...}. This is also required for compatibility with the translation module.
    - To mix text and equations, close the \text{...} block, write the equation, and then start a new \text{...} block for subsequent text.
    - If you need to include text within an equation (e.g., as a note in brackets), use slashes between words: e.g., (note\ between\ words).
    """

    def __init__(self) -> None:
        """Initialize the converter and optionally convert LaTeX to a Document."""
        self.template_docx: str | None = None

    @staticmethod
    def _apply_default_styles(doc: DocumentObject) -> None:
        """Apply default style settings to a new document.

        Args:
            doc: The Word Document object to apply styles to.
        """
        # Set default font for Normal style
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)

        # Set paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(8)
        paragraph_format.line_spacing = 1.15

        # Configure heading styles
        heading_configs = [
            ("Heading 1", 16, True, RGBColor(0, 40, 85)),
            ("Heading 2", 13, True, RGBColor(0, 40, 85)),
            ("Heading 3", 11, True, RGBColor(0, 40, 85)),
        ]

        for style_name, font_size, bold, color in heading_configs:
            if style_name in doc.styles:
                heading_style = doc.styles[style_name]
                heading_style.font.name = "Bebas Neue"
                heading_style.font.size = Pt(font_size)
                heading_style.font.bold = bold
                heading_style.font.color.rgb = color
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)

    @staticmethod
    def _formula(latex_string: str) -> BaseOxmlElement:
        """Convert a LaTeX equation string to an OMML XML element for Word.

        Args:
            latex_string: The LaTeX string representing the equation.
        """
        mathml_output = latex2mathml.converter.convert(latex_string)
        omml_output = mathml2omml.convert(mathml_output)
        xml_output = f'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml_output}</m:oMathPara>'
        return parse_xml(xml_output)[0]

    def convert_to_word(self, content: str, template_docx: str | None = None) -> DocumentObject:
        r"""
        Convert a LaTeX string (with text, equations, titles, sections, subsections, tables) to a Word Document object.
        Args:
            content: The LaTeX string, with lines separated by \newline, text in \text{...}, equations otherwise.
            template_docx: Optional path to an existing .docx file to use as a template. If provided, the content
                          will be added to this document. If None, a new document will be created.

        Returns
        -------
            The python-docx Document object (not saved).
        """
        self.template_docx = template_docx
        doc = Document(template_docx) if template_docx else Document()

        # Apply default styles if no template provided
        if template_docx is None:
            self._apply_default_styles(doc)

        matches = self._extract_structural_elements(content)
        parsed = self._build_parsed_content(content, matches)
        self._add_content_to_document(doc, parsed)
        return doc

    @staticmethod
    def _extract_structural_elements(content: str) -> list[dict[str, str | int]]:
        """Extract structural elements (titles, sections, tables) from LaTeX content.

        Args:
            content: The LaTeX string to parse.

        Returns
        -------
            A list of dictionaries containing type, content, start, and end positions.
        """
        # Patterns for all elements
        patterns: list[tuple[str, str, int]] = [
            ("title", r"\\title\{(.*?)\}", 0),
            ("section", r"\\section\{(.*?)\}", 0),
            ("subsection", r"\\subsection\{(.*?)\}", 0),
            ("subsubsection", r"\\subsubsection\{(.*?)\}", 0),
            ("table", r"\\begin\{table\}.*?\\end\{table\}", re.DOTALL),
        ]

        # Find all elements and their positions
        matches: list[dict[str, str | int]] = []
        for typ, pat, flags in patterns:
            matches.extend(
                {
                    "type": typ,
                    "content": m.group(1) if typ != "table" else m.group(0),
                    "start": m.start(),
                    "end": m.end(),
                }
                for m in re.finditer(pat, content, flags)
            )

        # Sort by position
        matches.sort(key=lambda x: int(x["start"]))
        return matches

    def _build_parsed_content(self, content: str, matches: list[dict[str, str | int]]) -> list[dict[str, str | bool]]:
        """Build parsed content list from structural elements and text/equations.

        Args:
            content: The original LaTeX string.
            matches: List of structural elements with their positions.

        Returns
        -------
            A list of parsed content items (text, equations, headings, tables).
        """
        parsed = []
        last_idx = 0

        for match in matches:
            # Text/equation between previous and current match
            start_pos = int(match["start"])
            end_pos = int(match["end"])
            if start_pos > last_idx:
                between = content[last_idx:start_pos]
                parsed.extend(self._parse_text_and_equations(between))
            parsed.append({"type": str(match["type"]), "content": str(match["content"])})
            last_idx = end_pos

        # Remaining text/equation after last match
        if last_idx < len(content):
            parsed.extend(self._parse_text_and_equations(content[last_idx:]))

        return parsed

    def _add_content_to_document(self, doc: DocumentObject, parsed: list[dict[str, str | bool]]) -> None:
        """Add parsed content items to the Word document.

        Args:
            doc: The Word Document object to add content to.
            parsed: List of parsed content items.
        """
        for item in parsed:
            item_type = str(item["type"])
            if item_type == "equation":
                self._add_equation(doc, str(item["content"]))
            elif item_type == "text":
                self._add_text(doc, item)
            elif item_type in ("title", "section", "subsection", "subsubsection"):
                self._add_heading(doc, item_type, str(item["content"]))
            elif item_type == "table":
                self._add_table_to_doc(doc, str(item["content"]))

    def _add_equation(self, doc: DocumentObject, content: str) -> None:
        """Add an equation to the document.

        Args:
            doc: The Word Document object.
            content: The LaTeX equation content.
        """
        p = doc.add_paragraph()
        p._p.append(self._formula(content))  # noqa: SLF001

    def _add_text(self, doc: DocumentObject, item: dict[str, str | bool]) -> None:
        """Add formatted text to the document."""
        para = doc.add_paragraph()
        run = para.add_run(str(item["content"]))
        if item.get("bold"):
            run.bold = True
        if item.get("italic"):
            run.italic = True

    def _add_heading(self, doc: DocumentObject, heading_type: str, content: str) -> None:
        """Add a heading to the document."""
        levels = {"title": 0, "section": 1, "subsection": 2, "subsubsection": 3}
        heading = doc.add_heading(content, level=levels[heading_type])

        # Apply Bebas Neue font if no template is used
        if self.template_docx is None:
            for run in heading.runs:
                run.font.name = "Bebas Neue"

    @staticmethod
    def _parse_text_and_equations(text: str) -> list[dict[str, str | bool]]:
        # Split by \newline for paragraph separation
        result: list[dict[str, str | bool]] = []

        # First, handle equation environments
        equation_pattern = r"\\begin\{equation\}(.*?)\\end\{equation\}"
        equation_matches = list(re.finditer(equation_pattern, text, re.DOTALL))

        if equation_matches:
            last_idx = 0
            for eq_match in equation_matches:
                # Process text before equation
                before_text = text[last_idx : eq_match.start()]
                if before_text.strip():
                    result.extend(LatexToWordConverter._parse_inline_content(before_text))

                # Process equation
                eq_content = eq_match.group(1).strip()

                # Extract tag if present
                tag_match = re.search(r"\\tag\{(.*?)\}", eq_content)

                # Remove \tag{...} from equation content
                if tag_match:
                    tag = tag_match.group(1)
                    eq_content = eq_content[: tag_match.start()] + eq_content[tag_match.end() :]
                    eq_content = eq_content.strip()
                    # Append tag as text to the right of the equation
                    eq_content += r"\text{   (" + tag + r")}"

                eq_item: dict[str, str | bool] = {"type": "equation", "content": eq_content}
                result.append(eq_item)

                last_idx = eq_match.end()

            # Process remaining text after last equation
            remaining = text[last_idx:]
            if remaining.strip():
                result.extend(LatexToWordConverter._parse_inline_content(remaining))
        else:
            # No equation environments, process as inline content
            result.extend(LatexToWordConverter._parse_inline_content(text))

        return result

    @staticmethod
    def _parse_inline_content(text: str) -> list[dict[str, str | bool]]:
        """Parse inline text and equations (not in equation environments)."""
        result: list[dict[str, str | bool]] = []
        for part in text.split(r"\newline"):
            # Pattern matches \textbf{...}, \textit{...}, \text{...}, or equation
            pattern = (
                r"\\textbf\{(.*?)\}"
                r"|\\textit\{(.*?)\}"
                r"|\\text\{(.*?)\}"
                r"|([^\n]+?)(?=(\\textbf\{|\\textit\{|\\text\{|$))"
            )
            for match in re.finditer(pattern, part):
                bold_content = match.group(1)
                italic_content = match.group(2)
                text_content = match.group(3)
                eq_content = match.group(4)
                if bold_content is not None:
                    result.append({"type": "text", "content": bold_content, "bold": True, "italic": False})
                elif italic_content is not None:
                    result.append({"type": "text", "content": italic_content, "bold": False, "italic": True})
                elif text_content is not None:
                    result.append({"type": "text", "content": text_content, "bold": False, "italic": False})
                elif eq_content is not None and eq_content.strip():
                    result.append({"type": "equation", "content": eq_content.strip()})
        return result

    @staticmethod
    def _parse_table_content(table_latex: str) -> list[list[str]]:
        """Parse LaTeX table content into rows and columns."""
        # Extract tabular content
        tabular_match = re.search(r"\\begin\{tabular\}\{.*?\}(.*?)\\end\{tabular\}", table_latex, re.DOTALL)
        if not tabular_match:
            return []

        tabular_content = tabular_match.group(1)

        # Remove LaTeX table commands
        tabular_content = re.sub(r"\\toprule|\\midrule|\\bottomrule|\\hline|\\centering", "", tabular_content)

        # Split by rows (\\)
        rows = re.split(r"\\\\", tabular_content)

        parsed_rows = []
        for row in rows:
            row_stripped = row.strip()
            if not row_stripped:
                continue
            # Split by columns (&)
            cells = row_stripped.split("&")
            parsed_cells = []
            for cell in cells:
                cell_stripped = cell.strip()
                # Extract text from \text{...} blocks
                text_matches = re.findall(r"\\text\{(.*?)\}", cell_stripped)
                cell_content = " ".join(text_matches) if text_matches else cell_stripped
                parsed_cells.append(cell_content)
            if parsed_cells:
                parsed_rows.append(parsed_cells)

        return parsed_rows

    def _add_table_to_doc(self, doc: DocumentObject, table_latex: str) -> None:
        """Parse LaTeX table and add it to the Word document."""
        rows_data = self._parse_table_content(table_latex)

        if not rows_data:
            doc.add_paragraph("[Empty table]")
            return

        # Create Word table
        num_rows = len(rows_data)
        num_cols = max(len(row) for row in rows_data) if rows_data else 0

        table = doc.add_table(rows=num_rows, cols=num_cols)
        if self.template_docx is None:
            table.style = "Medium Shading 1 Accent 5"

        # Fill table cells
        for i, row_data in enumerate(rows_data):
            for j, cell_content in enumerate(row_data):
                if j < num_cols:
                    table.rows[i].cells[j].text = cell_content
