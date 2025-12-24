"""Module to convert Report to a Word document."""

import re

import latex2mathml.converter
import mathml2omml
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import parse_xml
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


class ReportToWordConverter:
    r"""
    Converter to transform a Report (with text blocks, equations, titles, sections, tables) to a Word Document.

    Limitations and Usage Notes:
    - This converter currently supports Report input containing:
        - Text blocks using \text{...} (can also be \textbf{...} or \textit{...} for bold/italic).
        - Equations in \begin{equation}...\end{equation} environments (with optional \tag{...}).
        - Titles, sections, subsections, and subsubsections using \title{...}, \section{...}, \subsection{...}, \subsubsection{...}.
        - Tables in \begin{table}...\end{table} environments with tabular content
        - Inline equations (not in equation environments) between $...$.
        - New lines indicated by \newline.
        - Default styling for Normal text and Headings if no template is provided,
          or uses styles from a provided template (must be an empty .docx file featuring template styles for
          Normal, Title, Heading 1, Heading 2, and Heading 3).
    - The converter does NOT currently support:
        - Figures, images, or graphics.
    """

    def __init__(self) -> None:
        """Initialize the converter and optionally convert LaTeX to a Document."""
        self.template_docx: str | None = None

    def convert_to_word(self, content: str, template_docx: str | None = None) -> DocumentObject:
        r"""
        Convert a LaTeX string (with text, equations, titles, sections, subsections, tables) to a Word Document object.
        Args:
            content: The report string to convert.
            template_docx: Optional path to a .docx template file to use for styling.

        Returns
        -------
            The python-docx Document object (not saved).
        """
        self.template_docx = template_docx
        doc = Document(template_docx) if template_docx else Document()

        # Preprocess to remove LaTeX preamble commands
        content = self._preprocess_content(content)

        matches = self._extract_structural_elements(content)
        parsed = self._build_parsed_content(content, matches)
        self._add_content_to_document(doc, parsed)
        return doc

    @staticmethod
    def _preprocess_content(content: str) -> str:
        """Remove LaTeX preamble and document structure commands that should be ignored.

        Args:
            content: The raw LaTeX string.

        Returns
        -------
            Cleaned content with preamble commands removed
        """
        # Patterns to remove
        remove_patterns = [
            r"\\documentclass\{.*?\}",
            r"\\usepackage(\[.*?\])?\{.*?\}",
            r"\\geometry\{.*?\}",
            r"\\setstretch\{.*?\}",
            r"\\setlength\{.*?\}\{.*?\}",
            r"\\begin\{document\}",
            r"\\end\{document\}",
            r"\\date\{.*?\}",
            r"\\maketitle",
            r"\\begin{figure}.*?\\end{figure}",
            r"\\includegraphics(\[.*?\])?\{.*?\}",
        ]

        cleaned = content
        for pattern in remove_patterns:
            cleaned = re.sub(pattern, "", cleaned)

        return cleaned.strip()

    @staticmethod
    def _extract_structural_elements(content: str) -> list[dict[str, str | int]]:
        """Extract structural elements (titles, sections, tables) from LaTeX content.

        Args:
            content: The LaTeX string to parse.

        Returns
        -------
            A list of dictionaries containing type, content, start, and end positions.
        """
        # Patterns for all elements
        patterns: list[tuple[str, str, int]] = [
            ("title", r"\\title\{(.*?)\}", 0),
            ("section", r"\\section\{(.*?)\}", 0),
            ("subsection", r"\\subsection\{(.*?)\}", 0),
            ("subsubsection", r"\\subsubsection\{(.*?)\}", 0),
            ("table", r"\\begin\{table\}.*?\\end\{table\}", re.DOTALL),
        ]

        # Find all elements and their positions
        matches: list[dict[str, str | int]] = []
        for element_type, pattern, flags in patterns:
            matches.extend(
                {
                    "type": element_type,
                    "content": match.group(1) if element_type != "table" else match.group(0),
                    "start": match.start(),
                    "end": match.end(),
                }
                for match in re.finditer(pattern, content, flags)
            )

        # Sort by position
        matches.sort(key=lambda x: int(x["start"]))
        return matches

    def _build_parsed_content(self, content: str, matches: list[dict[str, str | int]]) -> list[dict[str, str | bool]]:
        """Build parsed content list from structural elements and text/equations.

        Args:
            content: The original LaTeX string.
            matches: List of structural elements with their positions.

        Returns
        -------
            A list of parsed content items (text, equations, headings, tables).
        """
        parsed = []
        last_idx = 0

        for match in matches:
            # Text/equation between previous and current match
            start_pos = int(match["start"])
            end_pos = int(match["end"])
            if start_pos > last_idx:
                between = content[last_idx:start_pos]
                parsed.extend(self._parse_text_and_equations(between))
            parsed.append({"type": str(match["type"]), "content": str(match["content"])})
            last_idx = end_pos

        # Remaining text/equation after last match
        if last_idx < len(content):
            parsed.extend(self._parse_text_and_equations(content[last_idx:]))

        return parsed

    def _add_content_to_document(self, doc: DocumentObject, parsed: list[dict[str, str | bool]]) -> None:
        """Add parsed content items to the Word document.

        Args:
            doc: The Word Document object to add content to.
            parsed: List of parsed content items.
        """
        for item in parsed:
            item_type = str(item["type"])
            if item_type == "equation":
                self._add_equation(doc, str(item["content"]))
            elif item_type == "text":
                self._add_text(doc, item)
            elif item_type in ("title", "section", "subsection", "subsubsection"):
                self._add_heading(doc, item_type, str(item["content"]))
            elif item_type == "table":
                self._add_table_to_doc(doc, str(item["content"]))

    def _add_heading(self, doc: DocumentObject, heading_type: str, content: str) -> None:
        """Add a heading to the document.

        Args:
            doc: The Word Document object.
            heading_type: Type of heading ('title', 'section', 'subsection', or 'subsubsection').
            content: The text content of the heading.
        """
        levels = {"title": 0, "section": 1, "subsection": 2, "subsubsection": 3}
        heading = doc.add_heading(content, level=levels[heading_type])

        # Apply default formatting if no template is provided
        if self.template_docx is None:
            heading_formats = {
                "title": (24, True, RGBColor(0, 40, 85)),
                "section": (18, True, RGBColor(0, 40, 85)),
                "subsection": (14, True, RGBColor(0, 40, 85)),
                "subsubsection": (14, False, RGBColor(0, 40, 85)),
            }
            font_size, bold, color = heading_formats[heading_type]
            for run in heading.runs:
                run.font.name = "Bebas Neue"
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run.font.color.rgb = color
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)

    def _add_text(self, doc: DocumentObject, item: dict[str, str | bool]) -> None:
        """Add formatted text (with possible inline equations) to the document.

        Args:
            doc: The Word Document object.
            item: Dictionary containing 'content' (text), and optional 'bold' and 'italic' flags.
        """
        para = doc.add_paragraph()
        content = str(item["content"])
        bold = bool(item.get("bold", False))
        italic = bool(item.get("italic", False))

        # Apply default formatting if no template is provided
        if self.template_docx is None:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.line_spacing = 1.15

        self._add_text_with_inline_math(para, content, bold=bold, italic=italic)

    def _add_text_with_inline_math(self, para: Paragraph, content: str, bold: bool = False, italic: bool = False) -> None:
        """Add text with inline math ($...$) to a paragraph, applying bold/italic as needed.

        Args:
            para: The paragraph object to add content to.
            content: The text content that may contain inline math ($...$).
            bold: Whether to apply bold formatting to text runs.
            italic: Whether to apply italic formatting to text runs.
        """
        # Split content by inline math ($...$)
        parts = re.split(r"(\$.*?\$)", content)
        for part in parts:
            if part.startswith("$") and part.endswith("$"):
                # Inline equation
                latex_eq = part[1:-1]
                para._p.append(self._formula(latex_eq.replace(r"\%", "%")))  # noqa: SLF001
            elif part:
                run = para.add_run(part.replace(r"\%", "%"))
                run.bold = bold
                run.italic = italic
                # Apply default formatting if no template is provided
                if self.template_docx is None:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_equation(self, doc: DocumentObject, content: str) -> None:
        r"""Add an equation to the document.

        If the equation is longer than 150 characters, split it on every '=' (except the first),
        or on every '\leq' or '\geq' (except the first), whichever comes first.

        Args:
            doc: The Word Document object.
            content: The LaTeX equation content.
        """
        if len(content) > 150:
            # Find all split points: '=', '\leq', '\geq'
            split_pattern = r"(=|\\leq|\\geq)"
            parts = re.split(split_pattern, content)
            # Find the first split point (if any)
            if len(parts) > 3:
                first = parts[0] + parts[1] + parts[2]
                rest = parts[3:]
                segments = [first]
                for i in range(0, len(rest), 2):
                    seg = "".join(rest[i : i + 2])
                    if seg:
                        segments.append(seg)
            else:
                segments = [content]
            for seg in segments:
                if seg.strip():
                    p = doc.add_paragraph()
                    p._p.append(self._formula(seg.strip()))  # noqa: SLF001
        else:
            p = doc.add_paragraph()
            p._p.append(self._formula(content))  # noqa: SLF001

    def _add_table_to_doc(self, doc: DocumentObject, table_latex: str) -> None:
        """Parse LaTeX table and add it to the Word document.

        Args:
            doc: The Word Document object.
            table_latex: LaTeX string containing a table environment.
        """
        rows_data = self._parse_table_content(table_latex)

        if not rows_data:
            doc.add_paragraph("[Empty table]")
            return

        # Create Word table
        num_rows = len(rows_data)
        num_cols = max(len(row) for row in rows_data) if rows_data else 0

        table = doc.add_table(rows=num_rows, cols=num_cols)

        # Apply default table style if no template is provided
        if self.template_docx is None:
            table.style = "Medium Shading 1 Accent 5" if "Medium Shading 1 Accent 5" in doc.styles else "Light Grid Accent 1"

        # Fill table cells (supporting inline math)
        for i, row_data in enumerate(rows_data):
            for j, cell_content in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    # Clear default paragraph and add content with inline math support
                    cell.text = ""
                    para = cell.paragraphs[0]
                    self._add_text_with_inline_math(para, cell_content)

    @staticmethod
    def _formula(latex_string: str) -> BaseOxmlElement:
        """Convert a LaTeX equation string (already in math mode) to an OMML XML element for Word.

        Args:
            latex_string: The LaTeX string representing the equation.

        Returns
        -------
            BaseOxmlElement containing the OMML representation of the equation.
        """
        mathml_output = latex2mathml.converter.convert(latex_string)
        omml_output = mathml2omml.convert(mathml_output)
        xml_output = f'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml_output}</m:oMathPara>'
        return parse_xml(xml_output)[0]

    @staticmethod
    def _parse_text_and_equations(text: str) -> list[dict[str, str | bool]]:
        r"""Parse text containing both inline text and equation environments.

        Args:
            text: LaTeX string that may contain \begin{equation}...\end{equation} blocks.

        Returns
        -------
            List of dictionaries representing text and equation items.
        """
        result: list[dict[str, str | bool]] = []

        # First, handle equation environments
        equation_pattern = r"\\begin\{equation\}(.*?)\\end\{equation\}"
        equation_matches = list(re.finditer(equation_pattern, text, re.DOTALL))

        if equation_matches:
            last_idx = 0
            for equation_match in equation_matches:
                # Process text before equation
                before_text = text[last_idx : equation_match.start()]
                if before_text.strip():
                    result.extend(ReportToWordConverter._parse_inline_content(before_text))

                # Process equation
                equation_content = equation_match.group(1).strip()

                # Extract tag if present
                tag_match = re.search(r"\\tag\{(.*?)\}", equation_content)

                # Remove \tag{...} from equation content
                if tag_match:
                    tag = tag_match.group(1)
                    equation_content = equation_content[: tag_match.start()] + equation_content[tag_match.end() :]
                    equation_content = equation_content.strip()
                    # Append tag as text to the right of the equation
                    equation_content += r"\text{   (" + tag + r")}"

                equation_item: dict[str, str | bool] = {"type": "equation", "content": equation_content}
                result.append(equation_item)

                last_idx = equation_match.end()

            # Process remaining text after last equation
            remaining = text[last_idx:]
            if remaining.strip():
                result.extend(ReportToWordConverter._parse_inline_content(remaining))
        else:
            # No equation environments, process as inline content
            result.extend(ReportToWordConverter._parse_inline_content(text))

        return result

    @staticmethod
    def _parse_inline_content(text: str) -> list[dict[str, str | bool]]:
        """Parse text (can contain inline text formula).

        Args:
            text: LaTeX string without equation environments.

        Returns
        -------
            List of dictionaries representing text and inline equation items.
        """
        result: list[dict[str, str | bool]] = []
        for part in text.split(r"\newline"):
            # Pattern matches \textbf{...}, \textit{...}, \text{...} with nested braces
            pattern = (
                r"\\textbf\{((?:[^{}]|\{[^{}]*\})*)\}"
                r"|\\textit\{((?:[^{}]|\{[^{}]*\})*)\}"
                r"|\\text\{((?:[^{}]|\{[^{}]*\})*)\}"
            )
            for match in re.finditer(pattern, part):
                bold_content = match.group(1)
                italic_content = match.group(2)
                text_content = match.group(3)
                if bold_content is not None:
                    result.append({"type": "text", "content": bold_content, "bold": True, "italic": False})
                elif italic_content is not None:
                    result.append({"type": "text", "content": italic_content, "bold": False, "italic": True})
                elif text_content is not None:
                    result.append({"type": "text", "content": text_content, "bold": False, "italic": False})
        return result

    @staticmethod
    def _parse_table_content(table_latex: str) -> list[list[str]]:
        """Parse LaTeX table content into rows and columns.

        Args:
            table_latex: LaTeX string containing a table environment.

        Returns
        -------
            List of rows, where each row is a list of cell contents.
        """
        # Extract tabular content
        tabular_match = re.search(r"\\begin\{tabular\}\{.*?\}(.*?)\\end\{tabular\}", table_latex, re.DOTALL)
        if not tabular_match:
            return []

        tabular_content = tabular_match.group(1)

        # Remove LaTeX table commands
        tabular_content = re.sub(r"\\toprule|\\midrule|\\bottomrule|\\hline|\\centering", "", tabular_content)

        # Split by rows (\\)
        rows = re.split(r"\\\\", tabular_content)

        parsed_rows = []
        for row in rows:
            row_stripped = row.strip()
            if not row_stripped:
                continue
            # Split by columns (&)
            cells = row_stripped.split("&")
            parsed_cells = []
            for cell in cells:
                cell_stripped = cell.strip()
                # Extract text from \text{...} blocks while preserving inline math
                # Handle nested braces: match \text{...} with up to one level of nested braces
                cell_content = re.sub(r"\\text\{((?:[^{}]|\{[^{}]*\})*)\}", r"\1", cell_stripped)
                parsed_cells.append(cell_content)
            if parsed_cells:
                parsed_rows.append(parsed_cells)

        return parsed_rows
